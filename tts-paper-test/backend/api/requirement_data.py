"""
CRUD API — 需求/方案/用例/评审 数据管理 + 文档上传解析 + 用例确认 + 用例导出
"""
import csv
import io
import json
import os
from pathlib import Path
from fastapi import APIRouter, HTTPException, Query, UploadFile, File
from fastapi.responses import StreamingResponse, PlainTextResponse
from pydantic import BaseModel, Field
from typing import Optional, Any
from services.storage_service import get_store

router = APIRouter(prefix="/api/v1/data", tags=["需求测试数据管理"])


# ====== 通用请求/响应模型 ======

class PaginatedResponse(BaseModel):
    items: list[dict[str, Any]]
    total: int
    page: int
    page_size: int
    total_pages: int


class CreateRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=200, description="标题")
    content: str = Field("", description="内容/Markdown")
    source: str = Field("manual", description="来源: manual / ai")
    status: str = Field("draft", description="状态")
    extra: dict[str, Any] = Field(default_factory=dict, description="扩展字段")


class UpdateRequest(BaseModel):
    title: Optional[str] = None
    content: Optional[str] = None
    source: Optional[str] = None
    status: Optional[str] = None
    case_status: Optional[str] = None
    extra: Optional[dict[str, Any]] = None


# ====== 文档上传解析 ======

UPLOAD_DIR = Path(__file__).resolve().parent.parent / "uploads"
UPLOAD_DIR.mkdir(exist_ok=True)

SUPPORTED_EXTENSIONS = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".pdf": "application/pdf",
}


def _parse_docx(path: str) -> str:
    """解析 .docx 文件"""
    from docx import Document
    doc = Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    # 也提取表格内容
    tables = []
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text for cell in row.cells]
            tables.append(" | ".join(cells))
    text = "\n".join(paragraphs)
    if tables:
        text += "\n\n--- 表格内容 ---\n" + "\n".join(tables)
    return text.strip()


def _parse_pdf(path: str) -> str:
    """解析 .pdf 文件"""
    import pdfplumber
    pages = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text()
            if text:
                pages.append(text)
    return "\n\n".join(pages).strip()


@router.post("/upload", summary="上传并解析文档")
async def upload_document(file: UploadFile = File(...)):
    """上传 .txt / .md / .docx / .pdf 文件，解析后返回文本内容"""
    ext = os.path.splitext(file.filename or "")[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise HTTPException(400, f"不支持的文件格式: {ext}，支持: {', '.join(SUPPORTED_EXTENSIONS.keys())}")

    # 保存文件
    file_path = UPLOAD_DIR / f"upload_{os.urandom(4).hex()}{ext}"
    content_bytes = await file.read()
    file_path.write_bytes(content_bytes)

    # 解析
    try:
        if ext in (".txt", ".md"):
            text = content_bytes.decode("utf-8", errors="replace").strip()
        elif ext == ".docx":
            text = _parse_docx(str(file_path))
        elif ext == ".pdf":
            text = _parse_pdf(str(file_path))
        else:
            text = ""
    except Exception as e:
        # 清理失败文件
        if file_path.exists():
            file_path.unlink()
        raise HTTPException(500, f"文件解析失败: {str(e)}")

    if not text:
        raise HTTPException(400, "文件内容为空或无法提取文本")

    return {
        "filename": file.filename,
        "content": text,
        "size": len(content_bytes),
        "chars": len(text),
    }


@router.get("/upload/supported", summary="获取支持的文件格式")
async def supported_formats():
    return {"formats": list(SUPPORTED_EXTENSIONS.keys())}


# ====== 用例确认（待确认→已确认） ======

class ConfirmCasesRequest(BaseModel):
    case_ids: list[str] = Field(..., min_length=1, description="要确认的用例ID列表")
    confirmed_by: str = Field("manual", description="确认方式: manual / ai")


@router.post("/cases/confirm", summary="批量确认用例（待确认→已确认）")
async def confirm_cases(body: ConfirmCasesRequest):
    """将指定的候选用例状态从 pending 改为 confirmed，同步到用例管理列表"""
    store = get_store("cases")
    confirmed = []
    for cid in body.case_ids:
        item = store.get(cid)
        if item and item.get("case_status") == "pending":
            updated = store.update(cid, {
                "case_status": "confirmed",
                "case_status_label": "已确认",
                "status": "completed",
                "confirmed_by": body.confirmed_by,
                "confirmed_at": __import__("datetime").datetime.now().isoformat(),
            })
            if updated:
                confirmed.append(updated)
    return {
        "ok": True,
        "confirmed": len(confirmed),
        "items": confirmed,
    }


@router.get("/cases/pending", summary="获取所有待确认用例")
async def list_pending_cases(pipeline_id: Optional[str] = None):
    """获取所有 case_status=pending 的用例，可选按 pipeline_id 筛选"""
    store = get_store("cases")
    result = store.list(page=1, page_size=1000, filters={"case_status": "pending"})
    items = result["items"]
    if pipeline_id:
        items = [i for i in items if i.get("pipeline_id") == pipeline_id]
    return {"items": items, "total": len(items)}


# ====== 用例导出 ======

@router.get("/cases/export", summary="导出用例（支持多种格式）")
async def export_cases(
    format: str = Query("md", description="导出格式: md / xlsx / csv / txt / xmind"),
    ids: Optional[str] = Query(None, description="指定用例ID（逗号分隔），不指定则导出全部"),
):
    """导出用例为多种格式"""
    store = get_store("cases")
    if ids:
        id_list = [x.strip() for x in ids.split(",") if x.strip()]
        all_items = []
        for iid in id_list:
            item = store.get(iid)
            if item:
                all_items.append(item)
    else:
        result = store.list(page=1, page_size=10000)
        all_items = result["items"]

    if not all_items:
        raise HTTPException(404, "没有可导出的用例")

    fmt = format.lower()

    if fmt == "md":
        return _export_md(all_items)
    elif fmt == "csv":
        return _export_csv(all_items)
    elif fmt == "txt":
        return _export_txt(all_items)
    elif fmt == "xlsx":
        return await _export_xlsx(all_items)
    elif fmt == "xmind":
        return await _export_xmind(all_items)
    else:
        raise HTTPException(400, f"不支持的导出格式: {fmt}，支持: md / xlsx / csv / txt / xmind")


def _export_md(items: list[dict]) -> PlainTextResponse:
    """导出为 Markdown"""
    lines = ["# 测试用例导出\n"]
    for i, item in enumerate(items, 1):
        title = item.get("title", "未命名")
        content = item.get("content", "")
        status = item.get("case_status_label") or item.get("status", "")
        source = item.get("source", "manual")
        lines.append(f"## {i}. {title}\n")
        lines.append(f"- **来源**: {source}")
        lines.append(f"- **状态**: {status}")
        lines.append(f"- **创建时间**: {item.get('created_at', '')}\n")
        if content:
            lines.append(f"{content}\n")
        lines.append("---\n")
    text = "\n".join(lines)
    return PlainTextResponse(text, media_type="text/markdown; charset=utf-8",
                             headers={"Content-Disposition": "attachment; filename=test-cases-export.md"})


def _export_csv(items: list[dict]) -> StreamingResponse:
    """导出为 CSV"""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["序号", "标题", "来源", "状态", "内容摘要", "创建时间"])
    for i, item in enumerate(items, 1):
        content = item.get("content", "")
        writer.writerow([
            i,
            item.get("title", "未命名"),
            item.get("source", "manual"),
            item.get("case_status_label") or item.get("status", ""),
            content[:200] + "..." if len(content) > 200 else content,
            item.get("created_at", ""),
        ])
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=test-cases-export.csv"},
    )


def _export_txt(items: list[dict]) -> PlainTextResponse:
    """导出为纯文本"""
    lines = ["测试用例导出", "=" * 40, ""]
    for i, item in enumerate(items, 1):
        title = item.get("title", "未命名")
        content = item.get("content", "")
        status = item.get("case_status_label") or item.get("status", "")
        lines.append(f"[{i}] {title}")
        lines.append(f"    来源: {item.get('source', 'manual')} | 状态: {status}")
        lines.append(f"    时间: {item.get('created_at', '')}")
        if content:
            # 只保留纯文本，去掉markdown标记
            import re
            plain = re.sub(r'[#*`\[\]]', '', content)
            lines.append(f"    内容: {plain[:300]}")
        lines.append("")
    text = "\n".join(lines)
    return PlainTextResponse(text, media_type="text/plain; charset=utf-8",
                             headers={"Content-Disposition": "attachment; filename=test-cases-export.txt"})


async def _export_xlsx(items: list[dict]) -> StreamingResponse:
    """导出为 Excel (.xlsx)"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment
    except ImportError:
        raise HTTPException(500, "导出xlsx需要安装 openpyxl: pip install openpyxl")

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "测试用例"

    # 表头
    headers = ["序号", "标题", "来源", "状态", "内容", "创建时间"]
    header_font = Font(bold=True, color="FFFFFF")
    header_fill = PatternFill(start_color="F59E0B", end_color="F59E0B", fill_type="solid")
    for col, h in enumerate(headers, 1):
        cell = ws.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = Alignment(horizontal="center")

    # 数据
    for i, item in enumerate(items, 2):
        ws.cell(row=i, column=1, value=i - 1)
        ws.cell(row=i, column=2, value=item.get("title", "未命名"))
        ws.cell(row=i, column=3, value=item.get("source", "manual"))
        ws.cell(row=i, column=4, value=item.get("case_status_label") or item.get("status", ""))
        ws.cell(row=i, column=5, value=item.get("content", "")[:500])
        ws.cell(row=i, column=6, value=item.get("created_at", ""))

    # 列宽
    ws.column_dimensions["A"].width = 8
    ws.column_dimensions["B"].width = 30
    ws.column_dimensions["C"].width = 10
    ws.column_dimensions["D"].width = 12
    ws.column_dimensions["E"].width = 60
    ws.column_dimensions["F"].width = 20

    output = io.BytesIO()
    wb.save(output)
    output.seek(0)
    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=test-cases-export.xlsx"},
    )


async def _export_xmind(items: list[dict]) -> PlainTextResponse:
    """导出为 XMind 兼容格式 (CSV结构化，可导入XMind)"""
    output = io.StringIO()
    output.write("测试用例导出\n")
    output.write("=" * 50 + "\n\n")
    for i, item in enumerate(items, 1):
        title = item.get("title", "未命名")
        output.write(f"## {i}. {title}\n")
        output.write(f"- 来源: {item.get('source', 'manual')}\n")
        output.write(f"- 状态: {item.get('case_status_label') or item.get('status', '')}\n")
        output.write(f"- 创建时间: {item.get('created_at', '')}\n\n")
    output.seek(0)
    return PlainTextResponse(
        output.getvalue(),
        media_type="text/plain; charset=utf-8",
        headers={"Content-Disposition": "attachment; filename=test-cases-export.xmind"},
    )


# ====== 通用CRUD (4个集合) ======
# 注意: 此路由通过 JSON 文件提供 CRUD 操作。
# 新代码应使用 api/test_lifecycle.py 中的 DB 版端点。
# JSON 存储将逐步废弃，数据可通过 scripts/migrate_json_to_db.py 迁移。

# DB模型映射（用于读操作，逐步替代JSON）
DB_MODEL_MAP = {}

# 延迟导入以避免循环依赖
def _get_db_model(collection: str):
    if not DB_MODEL_MAP:
        from models.test_lifecycle import Requirement, TestPlan, TestCase, Review
        DB_MODEL_MAP.update({
            "requirements": Requirement,
            "plans": TestPlan,
            "cases": TestCase,
            "reviews": Review,
        })
    return DB_MODEL_MAP.get(collection)


def _entity_router(collection: str, tag: str):
    """为指定集合生成 CRUD 路由"""
    store = get_store(collection)
    # DB 模型（如果有）用于读操作回退
    db_model_cls = _get_db_model(collection)

    @router.get(f"/{collection}", summary=f"查询{tag}列表")
    def list_items(
        page: int = Query(1, ge=1),
        page_size: int = Query(20, ge=1, le=100),
        search: str = Query("", description="全文搜索"),
        sort_by: str = Query("created_at"),
        sort_order: str = Query("desc", regex="^(asc|desc)$"),
        source: Optional[str] = Query(None, description="来源: manual / ai"),
        status: Optional[str] = Query(None, description="状态"),
        case_status: Optional[str] = Query(None, description="用例状态: pending / confirmed"),
    ) -> PaginatedResponse:
        filters = {}
        if source:
            filters["source"] = source
        if status:
            filters["status"] = status
        if case_status:
            filters["case_status"] = case_status
        return store.list(
            page=page, page_size=page_size,
            search=search, sort_by=sort_by,
            sort_order=sort_order, filters=filters,
        )

    @router.get(f"/{collection}/{{item_id}}", summary=f"获取{tag}详情")
    def get_item(item_id: str):
        item = store.get(item_id)
        if not item:
            raise HTTPException(404, f"{tag}不存在")
        return item

    @router.post(f"/{collection}", summary=f"创建{tag}")
    def create_item(body: CreateRequest):
        return store.create({
            "title": body.title,
            "content": body.content,
            "source": body.source,
            "status": body.status,
            **body.extra,
        })

    @router.put(f"/{collection}/{{item_id}}", summary=f"更新{tag}")
    def update_item(item_id: str, body: UpdateRequest):
        data = {k: v for k, v in body.model_dump(exclude_none=True).items() if v is not None}
        updated = store.update(item_id, data)
        if not updated:
            raise HTTPException(404, f"{tag}不存在")
        return updated

    @router.delete(f"/{collection}/{{item_id}}", summary=f"删除{tag}")
    def delete_item(item_id: str):
        ok = store.delete(item_id)
        if not ok:
            raise HTTPException(404, f"{tag}不存在")
        return {"ok": True}

    @router.post(f"/{collection}/batch-delete", summary=f"批量删除{tag}")
    def batch_delete(ids: list[str]):
        count = store.delete_many(ids)
        return {"ok": True, "deleted": count}


# 注册4个集合的路由
_entity_router("requirements", "需求")
_entity_router("plans", "方案")
_entity_router("cases", "用例")
_entity_router("reviews", "评审")


# ====== 流程记录（已废弃） ======
# 已废弃 — pipeline_service.py 已迁移到 DB。
# 前端已切换到 /api/v1/test-lifecycle/pipeline-records 端点。
# 保留仅用于向后兼容旧 JSON 数据。

@router.get("/pipeline-records", summary="查询流程记录列表")
def list_pipeline_records(
    page: int = Query(1, ge=1),
    page_size: int = Query(20, ge=1, le=100),
    search: str = Query("", description="搜索"),
):
    store = get_store("pipeline_records")
    # 列表不返回 stage_results（太大）
    result = store.list(page=page, page_size=page_size, search=search, sort_by="created_at", sort_order="desc")
    for item in result["items"]:
        item.pop("stage_results", None)
        item.pop("requirement_content", None)
    return result


@router.get("/pipeline-records/{record_id}", summary="获取流程记录详情（含所有阶段内容）")
def get_pipeline_record(record_id: str):
    store = get_store("pipeline_records")
    item = store.get(record_id)
    if not item:
        raise HTTPException(404, "流程记录不存在")
    return item
